"""格式适配层（v6 §5.1）。

PDF→pdfplumber（**表格优先，文本兜底**）· 扫描件→PaddleOCR（离线）·
XLSX/CSV→pandas · DOCX→python-docx · 图片→PaddleOCR。

**禁止使用 `pypdf`**（`CLAUDE.md` §11 反模式）—— 本模块只 import pdfplumber，
仓库里也没装 pypdf，装了也过不了单测里的那条断言。

「表格优先，文本兜底」的判据是**这一页有没有抽出表格**：抽到了就用表格，
抽不到才退回文本。不是「两个都要然后合并」—— 合并会让同一份数据出现两遍，
Diff 层立刻炸。
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pdfplumber

from backend.core.config import get_settings
from backend.core.errors import IngestionError
from backend.core.logging import get_logger
from backend.ingestion.repair import repair_cell, repair_text

logger = get_logger(__name__)

#: 一页里少于这么多字符就认为是扫描件，转 OCR
OCR_TRIGGER_CHARS = 16


@dataclass(frozen=True)
class ExtractedTable:
    """一张抽出来的表。`rows` 已逐格过修复层。"""

    page: int
    index: int
    rows: tuple[tuple[str, ...], ...]

    @property
    def header(self) -> tuple[str, ...]:
        return self.rows[0] if self.rows else ()

    @property
    def body(self) -> tuple[tuple[str, ...], ...]:
        return self.rows[1:] if self.rows else ()


@dataclass
class ExtractedDocument:
    """一份文档抽取后的中间形态。"""

    path: Path
    media_type: str
    #: 逐页文本，已过修复层
    pages: list[str] = field(default_factory=list)
    tables: list[ExtractedTable] = field(default_factory=list)
    #: 走了 OCR 的页码（1 起）
    ocr_pages: list[int] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(self.pages)

    @property
    def page_count(self) -> int:
        return len(self.pages)


def _clean_rows(raw_rows: list[list[str | None]]) -> tuple[tuple[str, ...], ...]:
    """逐格过修复层，并丢掉全空行。"""
    cleaned: list[tuple[str, ...]] = []
    for row in raw_rows:
        cells = tuple(repair_cell(cell) for cell in row)
        if any(cells):
            cleaned.append(cells)
    return tuple(cleaned)


def extract_pdf(path: Path, *, allow_ocr: bool = True) -> ExtractedDocument:
    """pdfplumber 抽取：表格优先，文本兜底，纯图片页转 OCR。"""
    doc = ExtractedDocument(path=path, media_type="application/pdf")
    with pdfplumber.open(str(path)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            raw_text = page.extract_text() or ""
            if len(raw_text.strip()) < OCR_TRIGGER_CHARS and allow_ocr:
                logger.info("PDF 页文本过少，转 OCR", path=str(path), page=page_no)
                raw_text = ocr_image_bytes(page.to_image(resolution=200).original, path, page_no)
                doc.ocr_pages.append(page_no)
            doc.pages.append(repair_text(raw_text))

            for table_no, table in enumerate(page.extract_tables(), start=1):
                rows = _clean_rows(table)
                if rows:
                    doc.tables.append(ExtractedTable(page=page_no, index=table_no, rows=rows))
    if not doc.pages:
        raise IngestionError(f"PDF 未抽出任何页：{path}", details={"path": str(path)})
    return doc


def ocr_image_bytes(image: Any, path: Path, page_no: int = 1) -> str:
    """PaddleOCR 离线识别。模型权重在 `.data/paddleocr`，不出网。"""
    ocr = _get_ocr_engine()
    import numpy as np

    result = ocr.predict(np.array(image.convert("RGB")))
    lines: list[str] = []
    for block in result or []:
        texts = block.get("rec_texts") if isinstance(block, dict) else None
        if texts:
            lines.extend(str(t) for t in texts)
    if not lines:
        logger.warning("OCR 未识别出文本", path=str(path), page=page_no)
    return "\n".join(lines)


_OCR_ENGINE: Any = None


def _get_ocr_engine() -> Any:
    """惰性加载 PaddleOCR —— 它启动要几秒，不用就不该付这个代价。"""
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        from paddleocr import PaddleOCR

        settings = get_settings()
        _OCR_ENGINE = PaddleOCR(
            lang="ch",
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            paddlex_config=None,
            device="cpu",
            **{"paddle_home": str(settings.PADDLEOCR_HOME)},
        )
    return _OCR_ENGINE


def extract_image(path: Path) -> ExtractedDocument:
    """图片走 OCR。"""
    from PIL import Image

    with Image.open(path) as img:
        text = ocr_image_bytes(img, path)
    return ExtractedDocument(
        path=path, media_type="image", pages=[repair_text(text)], ocr_pages=[1]
    )


def extract_tabular(path: Path, media_type: str) -> ExtractedDocument:
    """XLSX / CSV 走 pandas。每个 sheet 作为一张表。"""
    import pandas as pd

    doc = ExtractedDocument(path=path, media_type=media_type)
    if media_type == "text/csv":
        frames = {"csv": pd.read_csv(path, dtype=str, keep_default_na=False)}
    else:
        frames = pd.read_excel(path, sheet_name=None, dtype=str)

    for index, (sheet, frame) in enumerate(frames.items(), start=1):
        frame = frame.fillna("")
        header = tuple(repair_cell(str(c)) for c in frame.columns)
        body = [tuple(repair_cell(str(v)) for v in row) for row in frame.itertuples(index=False)]
        rows = (header, *body)
        doc.tables.append(ExtractedTable(page=1, index=index, rows=rows))
        doc.pages.append(repair_text(f"{sheet}\n" + "\n".join(" ".join(r) for r in rows)))
    return doc


def extract_docx(path: Path) -> ExtractedDocument:
    """DOCX 走 python-docx：段落进文本，表格进 tables。"""
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    doc = ExtractedDocument(
        path=path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    doc.pages.append(repair_text("\n".join(paragraphs)))
    for index, table in enumerate(document.tables, start=1):
        rows = tuple(tuple(repair_cell(cell.text) for cell in row.cells) for row in table.rows)
        rows = tuple(r for r in rows if any(r))
        if rows:
            doc.tables.append(ExtractedTable(page=1, index=index, rows=rows))
    return doc


#: 媒体类型 → 适配器
_DISPATCH = {
    "application/pdf": lambda p: extract_pdf(p),
    "image/png": extract_image,
    "image/jpeg": extract_image,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx,
}


def extract_document(path: Path, media_type: str) -> ExtractedDocument:
    """适配层主入口，按嗅探出的媒体类型分发。"""
    if media_type in _DISPATCH:
        return _DISPATCH[media_type](path)
    if media_type in (
        "text/csv",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ):
        return extract_tabular(path, media_type)
    raise IngestionError(
        f"没有适配 {media_type} 的抽取器",
        details={"media_type": media_type, "path": str(path)},
    )


__all__ = [
    "OCR_TRIGGER_CHARS",
    "ExtractedDocument",
    "ExtractedTable",
    "extract_document",
    "extract_docx",
    "extract_image",
    "extract_pdf",
    "extract_tabular",
    "ocr_image_bytes",
]
